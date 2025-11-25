from jinja2 import Environment, FileSystemLoader
try:
    from weasyprint import HTML, CSS
except OSError:
    print("WARNING: WeasyPrint system dependencies not found. PDF generation will fail.")
    HTML = None
    CSS = None
from typing import Dict, Optional
import os
from datetime import datetime
import uuid
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
try:
    from html4docx import HtmlToDocx
except ImportError:
    print("WARNING: html4docx not found or incompatible. DOCX generation will fail.")
    HtmlToDocx = None
import re
import base64
from bs4 import BeautifulSoup
import logging

logger = logging.getLogger(__name__)

class HTMLPDFGenerator:
    def __init__(self):
        template_dir = os.path.join(os.path.dirname(__file__), '../templates')
        self.env = Environment(loader=FileSystemLoader(template_dir))
        
        self.template_mapping = {
            'A': 'template_a_technical.html',
            'B': 'template_b_academic.html',
            'C': 'template_c_narrative.html',
            'D': 'template_d_business.html',
            'E': 'template_e_usa_support.html',
            'F': 'template_f_technical_testimony.html'
        }
    
    def assemble_letter(self, blocks: Dict[str, str], design: Dict, llm, custom_instructions: Optional[str] = None) -> str:
        """Assemble letter from blocks with AI refinement."""
        letter_content = "\n\n".join(blocks.values())
        return letter_content

    def _embed_logo_as_base64(self, logo_path: str) -> Optional[str]:
        """Convert logo to base64 data URI for embedding in PDF/HTML

        This prevents file:// URI issues and embeds logo directly in document.
        """
        if not logo_path or not os.path.exists(logo_path):
            return None

        try:
            with open(logo_path, 'rb') as f:
                logo_data = f.read()

            # Detect MIME type from file extension
            ext = os.path.splitext(logo_path)[1].lower()
            mime_types = {
                '.png': 'image/png',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.gif': 'image/gif',
                '.svg': 'image/svg+xml',
                '.webp': 'image/webp'
            }
            mime = mime_types.get(ext, 'image/png')

            # Convert to base64
            b64_data = base64.b64encode(logo_data).decode('utf-8')
            data_uri = f'data:{mime};base64,{b64_data}'

            logger.info(f"Logo embedded as base64 ({len(b64_data)} chars, {mime})")
            return data_uri

        except Exception as e:
            logger.error(f"Failed to embed logo as base64: {e}")
            return None

    def _validate_and_clean_html(self, html_content: str) -> str:
        """Validate and clean HTML content from LLM

        - Removes forbidden tags (html, head, body, script, style)
        - Fixes malformed HTML structure
        - Ensures valid nesting
        """
        try:
            soup = BeautifulSoup(html_content, 'html.parser')

            # Remove forbidden tags (keep content, remove wrapper)
            for tag_name in ['html', 'head', 'body', 'script', 'style']:
                for tag in soup.find_all(tag_name):
                    tag.unwrap()

            # Remove inline style attributes (we use CSS classes instead)
            for tag in soup.find_all(style=True):
                del tag['style']

            # Validate basic structure
            if not soup.find(['p', 'h1', 'h2', 'h3', 'div']):
                raise ValueError("No content elements found in HTML")

            cleaned_html = str(soup)
            logger.debug(f"HTML cleaned: {len(html_content)} → {len(cleaned_html)} chars")
            return cleaned_html

        except Exception as e:
            logger.error(f"HTML cleaning failed: {e}")
            raise ValueError(f"Invalid HTML structure: {e}")

    def _validate_html_quality(self, html: str, blocks: Dict[str, str]) -> dict:
        """Check if HTML meets quality standards

        Returns dict with 'valid', 'issues', and 'score' keys
        """
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text()
        issues = []

        # Check minimum content length
        word_count = len(text.split())
        if word_count < 1800:
            issues.append(f"Word count too low: {word_count} < 1800")
        elif word_count > 2800:
            issues.append(f"Word count too high: {word_count} > 2800")

        # Check for placeholder text
        placeholders = ['[INSERT]', '[TODO]', '[PLACEHOLDER]', 'Lorem ipsum', 'XXX', '###']
        for placeholder in placeholders:
            if placeholder.lower() in text.lower():
                issues.append(f"Placeholder text found: {placeholder}")

        # Check for forbidden terms
        forbidden = ['application', 'eb2-niw', 'eb2 niw', 'peticionário', 'visto', 'imigração']
        for term in forbidden:
            if term.lower() in text.lower():
                issues.append(f"Forbidden term found: {term}")

        # Check that we have actual content structure
        if not soup.find('p'):
            issues.append("No paragraphs found")

        # Calculate quality score (100 - 10 points per issue)
        score = max(0, 100 - (len(issues) * 10))

        return {
            'valid': len(issues) == 0,
            'issues': issues,
            'score': score,
            'word_count': word_count
        }


    def html_to_pdf(
        self, 
        html_content: str, 
        output_path: str, 
        design: Dict,
        logo_path: Optional[str] = None,
        recommender_info: Optional[Dict] = None
    ):
        """Convert HTML to PDF with logo and formatting - no templates"""
        
        # Embed logo as base64
        logo_data_uri = self._embed_logo_as_base64(logo_path) if logo_path else None
        
        # Build full HTML with logo header
        logo_html = ""
        if logo_data_uri:
            logo_html = f'<img src="{logo_data_uri}" style="height: 50px; margin-bottom: 20px;" />'
        
        full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            max-width: 8.5in;
            margin: 0.5in;
            padding: 0;
            color: #333;
        }}
        h2 {{
            font-size: 13pt;
            margin-top: 20px;
            margin-bottom: 10px;
            font-weight: bold;
            border-bottom: 1px solid #ddd;
            padding-bottom: 5px;
        }}
        p {{
            margin-bottom: 10px;
            text-align: justify;
        }}
        .header {{
            margin-bottom: 20px;
            border-bottom: 2px solid #333;
            padding-bottom: 10px;
        }}
        .signature {{
            margin-top: 30px;
        }}
    </style>
</head>
<body>
    <div class="header">
        {logo_html}
        <p><strong>Data:</strong> {datetime.now().strftime('%d de %B de %Y')}</p>
    </div>
    
    {html_content}
    
    <div class="signature">
        <p><strong>{recommender_info.get('name', 'Professional Recommender') if recommender_info else 'Professional Recommender'}</strong></p>
        <p>{recommender_info.get('title', '') if recommender_info else ''}</p>
        <p>{recommender_info.get('company', '') if recommender_info else ''}</p>
        <p>{recommender_info.get('location', '') if recommender_info else ''}</p>
    </div>
</body>
</html>"""
        
        # Create output directory
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Convert to PDF
        HTML(string=full_html).write_pdf(output_path)
        
        print(f"✅ PDF generated: {os.path.basename(output_path)}")

    def html_to_pdf_direct(self, complete_html: str, output_path: str):
        """
        Convert complete HTML document to PDF (no additional wrapping).
        Used when HTML is already a full document from HTMLDesigner.

        Args:
            complete_html: Complete HTML document string (DOCTYPE to </html>)
            output_path: Path for output PDF file
        """
        if HTML is None:
            raise RuntimeError("WeasyPrint not available - cannot generate PDF")

        # Create output directory
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        # Convert to PDF directly
        try:
            HTML(string=complete_html).write_pdf(output_path)
            logger.info(f"PDF generated: {os.path.basename(output_path)}")
            print(f"✅ PDF generated: {os.path.basename(output_path)}")
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            raise ValueError(f"Failed to generate PDF: {e}")

    def _process_html_element_to_docx(self, element, doc, paragraph=None):
        """Recursively process HTML element and add to DOCX document with formatting preservation"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        if isinstance(element, str):
            # Text node
            if paragraph:
                run = paragraph.add_run(element)
                return run
            else:
                p = doc.add_paragraph(element)
                return p

        tag = element.name if hasattr(element, 'name') else None

        if tag == 'p':
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for child in element.children:
                self._process_html_element_to_docx(child, doc, p)
            return p

        elif tag in ['h1', 'h2', 'h3'] and tag:
            level = int(tag[1]) if tag and len(tag) > 1 else 1
            text = element.get_text()
            p = doc.add_heading(text, level=level)
            return p

        elif tag in ['strong', 'b']:
            if paragraph:
                text = element.get_text()
                run = paragraph.add_run(text)
                run.bold = True
                return run

        elif tag in ['em', 'i']:
            if paragraph:
                text = element.get_text()
                run = paragraph.add_run(text)
                run.italic = True
                return run

        elif tag in ['ul', 'ol']:
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet' if tag == 'ul' else 'List Number')
                for child in li.children:
                    self._process_html_element_to_docx(child, doc, p)

        elif tag == 'table':
            rows = element.find_all('tr')
            if not rows:
                return None

            # Count columns
            first_row = rows[0]
            cols = len(first_row.find_all(['th', 'td']))

            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = 'Light Grid Accent 1'

            for i, row in enumerate(rows):
                cells = row.find_all(['th', 'td'])
                for j, cell in enumerate(cells):
                    table.rows[i].cells[j].text = cell.get_text()
                    # Bold header cells
                    if cell.name == 'th':
                        for paragraph in table.rows[i].cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

            return table

        elif tag == 'br':
            if paragraph:
                paragraph.add_run('\n')

        elif tag == 'blockquote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            for child in element.children:
                self._process_html_element_to_docx(child, doc, p)
            return p

        elif tag == 'div':
            # Process div children
            for child in element.children:
                self._process_html_element_to_docx(child, doc)

        else:
            # Unknown tag - process children
            if hasattr(element, 'children'):
                for child in element.children:
                    self._process_html_element_to_docx(child, doc, paragraph)

        return None

    def html_to_docx(
        self,
        html_content: str,
        output_path: str,
        design: Dict,
        logo_path: Optional[str] = None,
        recommender_info: Optional[Dict] = None
    ):
        """Convert HTML to editable DOCX with improved formatting preservation

        CRITICAL FIX: Uses custom element-by-element conversion instead of html4docx
        This preserves bold, italic, tables, lists, and all formatting.
        """

        # Create new document
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Add logo at top if available
        if logo_path and os.path.exists(logo_path):
            try:
                doc.add_picture(logo_path, width=Inches(2.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            except Exception as e:
                print(f"⚠️ Could not add logo to DOCX: {e}")

        # Add header info
        if recommender_info:
            header_p = doc.add_paragraph()
            header_run = header_p.add_run(f"{recommender_info.get('name', '')}\n")
            header_run.bold = True
            header_run.font.size = Pt(12)

            if recommender_info.get('title'):
                header_p.add_run(f"{recommender_info['title']}\n").font.size = Pt(10)
            if recommender_info.get('company'):
                header_p.add_run(f"{recommender_info['company']}\n").font.size = Pt(10)
            if recommender_info.get('location'):
                header_p.add_run(f"{recommender_info['location']}\n").font.size = Pt(10)

        # Add date
        date_p = doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
        date_p.add_run('\n')

        # CRITICAL FIX: Custom HTML to DOCX conversion preserving formatting
        try:
            soup = BeautifulSoup(html_content, 'html.parser')

            # Process top-level elements
            for element in soup.children:
                if hasattr(element, 'name') and element.name:  # Skip text nodes at root level
                    self._process_html_element_to_docx(element, doc)

            print(f"✅ HTML converted to DOCX with formatting preservation")

        except Exception as e:
            logger.error(f"HTML to DOCX conversion failed: {e}")
            print(f"⚠️ HTML parsing error, using improved fallback: {e}")

            # Improved fallback: Parse and add with basic formatting
            soup = BeautifulSoup(html_content, 'html.parser')
            for p_tag in soup.find_all('p'):
                text = p_tag.get_text()
                if text.strip():
                    doc.add_paragraph(text)

        # Create output directory if needed
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        # Save DOCX
        doc.save(output_path)

        style_id = design.get('unique_id', 'STYLE_DEFAULT')
        print(f"✅ Editable DOCX generated with style {style_id}: {os.path.basename(output_path)}")

    def html_to_docx_direct(self, complete_html: str, output_path: str):
        """
        Convert complete HTML document to DOCX (extracts body content).
        Used when HTML is already a full document from HTMLDesigner.

        Args:
            complete_html: Complete HTML document string (DOCTYPE to </html>)
            output_path: Path for output DOCX file
        """
        try:
            # Parse HTML to extract body content
            soup = BeautifulSoup(complete_html, 'html.parser')
            body = soup.find('body')

            if not body:
                raise ValueError("No <body> tag found in HTML")

            # Create new DOCX document
            doc = Document()

            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Process all elements in body
            for element in body.children:
                if hasattr(element, 'name') and element.name:
                    self._process_html_element_to_docx(element, doc)

            # Create output directory if needed
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            # Save DOCX
            doc.save(output_path)

            logger.info(f"DOCX generated: {os.path.basename(output_path)}")
            print(f"✅ Editable DOCX generated: {os.path.basename(output_path)}")

        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")
            raise ValueError(f"Failed to generate DOCX: {e}")

            doc = Document()
            soup = BeautifulSoup(complete_html, 'html.parser')
            
            # Try to extract page margins from CSS if possible, otherwise default
            # (Simplification: just use standard margins)
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Extract Body Content
            body = soup.find('body')
            if not body:
                raise ValueError("No body tag found")

            # Simple recursive parser to handle basic formatting
            def process_element(element, paragraph=None):
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    p = doc.add_paragraph(style=f'Heading {element.name[1]}')
                    p.add_run(element.get_text().strip())
                elif element.name == 'p':
                    p = doc.add_paragraph()
                    # Process children for bold/italic
                    for child in element.children:
                        if child.name == 'strong' or child.name == 'b':
                            p.add_run(child.get_text()).bold = True
                        elif child.name == 'em' or child.name == 'i':
                            p.add_run(child.get_text()).italic = True
                        elif isinstance(child, str):
                            p.add_run(child)
                elif element.name == 'ul':
                    for li in element.find_all('li', recursive=False):
                        doc.add_paragraph(li.get_text().strip(), style='List Bullet')
                elif element.name == 'ol':
                    for li in element.find_all('li', recursive=False):
                        doc.add_paragraph(li.get_text().strip(), style='List Number')
                elif element.name == 'div':
                    for child in element.children:
                        if child.name: process_element(child)
                # Handle other block elements...

            # Start processing
            for child in body.children:
                if child.name:
                    process_element(child)

            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            print(f"✅ DOCX generated directly: {os.path.basename(output_path)}")

        except Exception as e:
            print(f"⚠️ DOCX direct generation failed, using fallback: {e}")
            # Fallback to simple text extraction
            self.html_to_docx(complete_html, output_path, {}, None, None)
# Keep backward compatibility
class DOCXGenerator(HTMLPDFGenerator):
    """Backward compatibility wrapper - now generates PDFs with HTML templates"""
    
    def markdown_to_docx(self, markdown_content: str, output_path: str, design: Dict, logo_path: Optional[str] = None):
        """Legacy method - redirects to HTML/PDF generation"""
        # Extract recommender info from markdown if present
        recommender_info = {
            'name': design.get('assigned_recommender', 'Professional Recommender'),
            'title': '',
            'company': '',
            'location': ''
        }
        
        # Change extension from .docx to .pdf
        if output_path.endswith('.docx'):
            output_path = output_path.replace('.docx', '.pdf')
        
        self.html_to_pdf(markdown_content, output_path, design, logo_path, recommender_info)
