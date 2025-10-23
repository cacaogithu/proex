from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Dict, Optional
import os
import re


class DOCXGenerator:
    def assemble_letter(self, blocks: Dict[str, str], design: Dict, llm) -> str:
        """Use Claude 4.5 Sonnet for premium HTML/document assembly"""
        combined_blocks = f"""
# BLOCO 3
{blocks.get('block3', '')}

# BLOCO 4
{blocks.get('block4', '')}

# BLOCO 5
{blocks.get('block5', '')}

# BLOCO 6
{blocks.get('block6', '')}

# BLOCO 7
{blocks.get('block7', '')}
"""
        
        prompt = f"""# ROLE
Você é um revisor de classe mundial especializado em cartas de recomendação profissionais. 
Receba 5 blocos de uma carta de testemunho e produza o texto completo em formato EDITÁVEL.

**PERSONA**: {design.get('tone_instructions', '')}

# INPUTS
{combined_blocks}

# INSTRUÇÕES CRÍTICAS
1. Leia todos os blocos e crie uma narrativa coesa
2. Verifique transições suaves entre seções
3. REMOVA palavras problemáticas: "inferência lógica", "inferência técnica", "nexo causal" se mal usadas
4. NUNCA mencione: "application", "EB2-NIW", "peticionário", "visto", "imigração"
5. Seja autêntico e pessoal (primeira pessoa)
6. Output: APENAS texto em Markdown limpo (sem JSON, YAML, ou code blocks)
7. Adicione cabeçalho profissional com:
   - Nome do recomendador: {design.get('assigned_recommender', '')}
   - Cargo (se disponível nos blocos)
   - Data atual

# ESTRUTURA DO DOCUMENTO
1. Cabeçalho (nome, cargo, data)
2. Saudação formal ("A quem possa interessar," ou similar)
3. Blocos integrados em narrativa fluida
4. Encerramento formal
5. Assinatura (nome do recomendador)

# HETEROGENEIDADE
- Garanta que este testemunho tenha voz única
- Evite clichês e repetições
- Use causalidade direta: "Realizou X, gerando Y resultado"
- Mantenha tom profissional mas humano

# FORMATAÇÃO PARA WORD
- Use **negrito** para ênfases importantes
- Use itálico para *citações* ou termos técnicos
- Quebre parágrafos adequadamente (não muito longos)
- Adicione espaços entre seções

# TODO EM PORTUGUÊS BRASILEIRO
"""
        
        import time
        max_retries = 3
        for attempt in range(max_retries):
            try:
                # Using Claude 4.5 Sonnet - best quality for document assembly
                response = llm.client.chat.completions.create(
                    model=llm.models["premium"],
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.7
                )
                
                return response.choices[0].message.content
            except Exception as e:
                if "429" in str(e) and attempt < max_retries - 1:
                    wait_time = (2 ** attempt) * 3
                    print(f"⏳ Rate limit hit, waiting {wait_time}s before retry {attempt + 1}/{max_retries}...")
                    time.sleep(wait_time)
                    continue
                if attempt == max_retries - 1:
                    print(f"Error assembling letter: {str(e)}")
        
        return combined_blocks
    
    def markdown_to_docx(
        self, 
        markdown_text: str, 
        output_path: str, 
        design: Dict,
        logo_path: Optional[str] = None
    ):
        """Convert markdown to editable DOCX with logo"""
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        doc = Document()
        
        # Add logo if available
        if logo_path and os.path.exists(logo_path):
            try:
                # Add logo aligned to right
                logo_paragraph = doc.add_paragraph()
                logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                run = logo_paragraph.add_run()
                run.add_picture(logo_path, width=Inches(1.5))
            except Exception as e:
                print(f"⚠️ Could not add logo: {str(e)}")
        
        # Parse markdown and convert to DOCX
        lines = markdown_text.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Empty line - add spacing
                doc.add_paragraph()
                continue
            
            # Headers
            if line.startswith('# '):
                heading = doc.add_heading(line[2:], level=1)
                heading.runs[0].font.size = Pt(16)
                heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            elif line.startswith('## '):
                heading = doc.add_heading(line[3:], level=2)
                heading.runs[0].font.size = Pt(14)
            elif line.startswith('### '):
                heading = doc.add_heading(line[4:], level=3)
                heading.runs[0].font.size = Pt(12)
            
            # Lists
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.5)
            elif re.match(r'^\d+\. ', line):
                content = re.sub(r'^\d+\. ', '', line)
                p = doc.add_paragraph(content, style='List Number')
                p.paragraph_format.left_indent = Inches(0.5)
            
            # Regular paragraphs
            else:
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)
        
        # Save document
        try:
            doc.save(output_path)
            print(f"✓ DOCX generated successfully: {output_path}")
        except Exception as e:
            print(f"✗ Error generating DOCX: {str(e)}")
            raise
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add text with markdown formatting (bold, italic) to paragraph"""
        # Split by bold and italic markers
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|__[^_]+__|_[^_]+_)', text)
        
        for part in parts:
            if not part:
                continue
            
            run = paragraph.add_run()
            
            # Bold
            if part.startswith('**') and part.endswith('**'):
                run.text = part[2:-2]
                run.bold = True
            elif part.startswith('__') and part.endswith('__'):
                run.text = part[2:-2]
                run.bold = True
            
            # Italic
            elif part.startswith('*') and part.endswith('*'):
                run.text = part[1:-1]
                run.italic = True
            elif part.startswith('_') and part.endswith('_'):
                run.text = part[1:-1]
                run.italic = True
            
            # Regular text
            else:
                run.text = part
            
            # Set font
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
