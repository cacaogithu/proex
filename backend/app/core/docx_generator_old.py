from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Dict, Optional
import os
import re
import markdown as md
from bs4 import BeautifulSoup
from bs4.element import NavigableString


class DOCXGenerator:
    def assemble_letter(self, blocks: Dict[str, str], design: Dict, llm) -> str:
        """Use Claude 4.5 Sonnet for premium HTML/document assembly"""
        combined_blocks = f"""
# BLOCO 3
{blocks.get('block3', '')}

# BLOCO 4
{blocks.get('block4', '')}

# BLOCO 5
{blocks.get('block5', '')}

# BLOCO 6
{blocks.get('block6', '')}

# BLOCO 7
{blocks.get('block7', '')}
"""
        
        prompt = f"""# ROLE
Você é um revisor de classe mundial especializado em cartas de recomendação profissionais. 
Receba 5 blocos de uma carta de testemunho e produza o texto completo em formato EDITÁVEL.

**PERSONA**: {design.get('tone_instructions', '')}

# INPUTS
{combined_blocks}

# INSTRUÇÕES CRÍTICAS
1. Leia todos os blocos e crie uma narrativa coesa
2. Verifique transições suaves entre seções
3. REMOVA palavras problemáticas: "inferência lógica", "inferência técnica", "nexo causal" se mal usadas
4. NUNCA mencione: "application", "EB2-NIW", "peticionário", "visto", "imigração"
5. Seja autêntico e pessoal (primeira pessoa)
6. Output: APENAS texto em Markdown limpo (sem JSON, YAML, ou code blocks)
7. Adicione cabeçalho profissional com:
   - Nome do recomendador: {design.get('assigned_recommender', '')}
   - Cargo (se disponível nos blocos)
   - Data atual

# ESTRUTURA DO DOCUMENTO
1. Cabeçalho (nome, cargo, data)
2. Saudação formal ("A quem possa interessar," ou similar)
3. Blocos integrados em narrativa fluida
4. Encerramento formal
5. Assinatura (nome do recomendador)

# HETEROGENEIDADE
- Garanta que este testemunho tenha voz única
- Evite clichês e repetições
- Use causalidade direta: "Realizou X, gerando Y resultado"
- Mantenha tom profissional mas humano

# FORMATAÇÃO PARA WORD
- Use **negrito** para ênfases importantes
- Use itálico para *citações* ou termos técnicos
- Quebre parágrafos adequadamente (não muito longos)
- Adicione espaços entre seções

# TODO EM PORTUGUÊS BRASILEIRO
"""
        
        import time
        max_retries = 3
        for attempt in range(max_retries):
            try:
                # Using Claude 4.5 Sonnet - best quality for document assembly
                response = llm.client.chat.completions.create(
                    model=llm.models["premium"],
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.7
                )
                
                return response.choices[0].message.content
            except Exception as e:
                if "429" in str(e) and attempt < max_retries - 1:
                    wait_time = (2 ** attempt) * 3
                    print(f"⏳ Rate limit hit, waiting {wait_time}s before retry {attempt + 1}/{max_retries}...")
                    time.sleep(wait_time)
                    continue
                if attempt == max_retries - 1:
                    print(f"Error assembling letter: {str(e)}")
        
        return combined_blocks
    
    def markdown_to_docx(
        self, 
        markdown_text: str, 
        output_path: str, 
        design: Dict,
        logo_path: Optional[str] = None
    ):
        """Convert markdown to editable DOCX with logo using proper markdown parser"""
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        doc = Document()
        
        # Add logo if available
        if logo_path and os.path.exists(logo_path):
            try:
                # Add logo aligned to right
                logo_paragraph = doc.add_paragraph()
                logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                run = logo_paragraph.add_run()
                run.add_picture(logo_path, width=Inches(1.5))
                doc.add_paragraph()  # Spacing after logo
            except Exception as e:
                print(f"⚠️ Could not add logo: {str(e)}")
        
        # Convert markdown to HTML first for better parsing
        html_content = md.markdown(markdown_text, extensions=['extra', 'nl2br', 'tables'])
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Parse HTML and convert to DOCX
        self._parse_html_to_docx(soup, doc)
        
        # Save document
        try:
            doc.save(output_path)
            print(f"✓ DOCX generated successfully: {output_path}")
        except Exception as e:
            print(f"✗ Error generating DOCX: {str(e)}")
            raise
    
    def _parse_html_to_docx(self, soup, doc):
        """Parse HTML elements and add to DOCX document"""
        for element in soup.children:
            if element.name is None:
                # Text node
                continue
            
            # Headings
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                heading = doc.add_heading(element.get_text(), level=level)
                if level == 1:
                    heading.runs[0].font.size = Pt(16)
                    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
                elif level == 2:
                    heading.runs[0].font.size = Pt(14)
                elif level == 3:
                    heading.runs[0].font.size = Pt(12)
            
            # Paragraphs
            elif element.name == 'p':
                p = doc.add_paragraph()
                self._add_html_formatted_text(p, element)
            
            # Lists
            elif element.name == 'ul':
                self._parse_list(doc, element, 'List Bullet', 0)
            
            elif element.name == 'ol':
                self._parse_list(doc, element, 'List Number', 0)
            
            # Tables
            elif element.name == 'table':
                rows = element.find_all('tr')
                if rows:
                    cols = len(rows[0].find_all(['th', 'td']))
                    table = doc.add_table(rows=len(rows), cols=cols)
                    table.style = 'Light Grid Accent 1'
                    
                    for i, row in enumerate(rows):
                        cells = row.find_all(['th', 'td'])
                        for j, cell in enumerate(cells):
                            table.rows[i].cells[j].text = cell.get_text(strip=True)
            
            # Blockquote
            elif element.name == 'blockquote':
                p = doc.add_paragraph(element.get_text())
                p.style = 'Intense Quote'
            
            # Line break
            elif element.name == 'br':
                doc.add_paragraph()
    
    def _parse_list(self, doc, list_element, style, level):
        """Parse list recursively to handle nested lists"""
        for li in list_element.find_all('li', recursive=False):
            # Check for nested lists
            nested_ul = li.find('ul', recursive=False)
            nested_ol = li.find('ol', recursive=False)
            
            # Get text without nested list content
            li_text = li.get_text(strip=True)
            if nested_ul:
                nested_text = nested_ul.get_text(strip=True)
                li_text = li_text.replace(nested_text, '').strip()
            if nested_ol:
                nested_text = nested_ol.get_text(strip=True)
                li_text = li_text.replace(nested_text, '').strip()
            
            # Add list item
            p = doc.add_paragraph(style=style)
            p.paragraph_format.left_indent = Inches(0.5 + (level * 0.3))
            self._add_html_formatted_text(p, li)
            
            # Process nested lists
            if nested_ul:
                self._parse_list(doc, nested_ul, 'List Bullet', level + 1)
            if nested_ol:
                self._parse_list(doc, nested_ol, 'List Number', level + 1)
    
    def _add_html_formatted_text(self, paragraph, element):
        """Add text with HTML formatting (bold, italic, links) to paragraph"""
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                # Only skip completely empty strings, preserve whitespace
                if text:
                    # Check if we're inside a link
                    if element.name == 'a' and element.get('href'):
                        self._add_hyperlink(paragraph, text, element['href'])
                    else:
                        run = paragraph.add_run(text)
                        self._apply_formatting(run, element)
            elif child.name == 'br':
                # Add line break
                paragraph.add_run().add_break()
            elif child.name == 'a' and child.get('href'):
                # Handle hyperlinks
                link_text = child.get_text()
                self._add_hyperlink(paragraph, link_text, child['href'])
            elif child.name in ['strong', 'b', 'em', 'i', 'code', 'p', 'span', 'div']:
                # Recursively handle formatted text and block elements
                self._add_html_formatted_text(paragraph, child)
            elif child.name in ['ul', 'ol']:
                # Skip nested lists (handled separately)
                continue
    
    def _apply_formatting(self, run, element):
        """Apply formatting based on HTML element"""
        # Check parent tags for formatting
        parent = element
        while parent:
            if parent.name in ['strong', 'b']:
                run.bold = True
            if parent.name in ['em', 'i']:
                run.italic = True
            if parent.name == 'code':
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            parent = parent.parent if hasattr(parent, 'parent') else None
        
        # Set default font
        if not run.font.name:
            run.font.name = 'Calibri'
        if not run.font.size:
            run.font.size = Pt(11)
    
    def _add_hyperlink(self, paragraph, text, url):
        """Add a hyperlink to the paragraph"""
        # For simplicity, just add as blue underlined text with URL in parentheses
        # Full hyperlink implementation requires manipulating XML directly
        run = paragraph.add_run(f"{text} ")
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.font.underline = True
        
        # Add URL in parentheses
        url_run = paragraph.add_run(f"({url})")
        url_run.font.size = Pt(9)
        url_run.font.color.rgb = RGBColor(100, 100, 100)
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add text with markdown formatting (bold, italic) to paragraph"""
        # Split by bold and italic markers
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|__[^_]+__|_[^_]+_)', text)
        
        for part in parts:
            if not part:
                continue
            
            run = paragraph.add_run()
            
            # Bold
            if part.startswith('**') and part.endswith('**'):
                run.text = part[2:-2]
                run.bold = True
            elif part.startswith('__') and part.endswith('__'):
                run.text = part[2:-2]
                run.bold = True
            
            # Italic
            elif part.startswith('*') and part.endswith('*'):
                run.text = part[1:-1]
                run.italic = True
            elif part.startswith('_') and part.endswith('_'):
                run.text = part[1:-1]
                run.italic = True
            
            # Regular text
            else:
                run.text = part
            
            # Set font
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
