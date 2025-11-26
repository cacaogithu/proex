import pdfplumber
from docx import Document
from typing import Dict, Any
import os

# Configuration
STORAGE_BASE_DIR = os.getenv('STORAGE_BASE_DIR', 'backend/storage')


class PDFExtractor:
    def extract_text(self, file_path: str) -> str:
        """Extract text from PDF or DOCX files"""
        if file_path.endswith('.docx'):
            return self._extract_docx(file_path)
        else:
            return self._extract_pdf(file_path)
    
    def _extract_pdf(self, pdf_path: str) -> str:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                return text.strip()
        except Exception as e:
            print(f"Error extracting PDF {pdf_path}: {str(e)}")
            return ""
    
    def _extract_docx(self, docx_path: str) -> str:
        try:
            doc = Document(docx_path)
            text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error extracting DOCX {docx_path}: {str(e)}")
            return ""

    def extract_all_files(self, submission_id: str) -> Dict[str, Any]:
        base_path = os.path.join(STORAGE_BASE_DIR, "uploads", submission_id)
        
        extracted = {}
        
        quadro_path = f"{base_path}/quadro.pdf"
        if os.path.exists(quadro_path):
            extracted["quadro"] = self.extract_text(quadro_path)
        else:
            extracted["quadro"] = ""
        
        cv_path = f"{base_path}/cv.pdf"
        if os.path.exists(cv_path):
            extracted["cv"] = self.extract_text(cv_path)
        else:
            extracted["cv"] = ""
        
        estrategia_path = f"{base_path}/estrategia.pdf"
        if os.path.exists(estrategia_path):
            extracted["estrategia"] = self.extract_text(estrategia_path)
        else:
            extracted["estrategia"] = ""
        
        onenote_path = f"{base_path}/onenote.pdf"
        if os.path.exists(onenote_path):
            extracted["onenote"] = self.extract_text(onenote_path)
        else:
            extracted["onenote"] = ""
        
        # Extract additional attached documents
        additional_docs = []
        i = 0
        while True:
            doc_path = f"{base_path}/attached_{i}.pdf"
            if os.path.exists(doc_path):
                additional_docs.append({
                    'filename': f'attached_{i}.pdf',
                    'text': self.extract_text(doc_path)
                })
                i += 1
            else:
                break
        extracted["additional_documents"] = additional_docs
        
        testimonials = []
        i = 0
        while True:
            testimonial_path = f"{base_path}/testimonial_{i}.pdf"
            if os.path.exists(testimonial_path):
                testimonials.append(self.extract_text(testimonial_path))
                i += 1
            else:
                break
        
        extracted["testimonials"] = testimonials
        
        return extracted
